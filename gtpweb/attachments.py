from __future__ import annotations

import base64
import io
import mimetypes
import re
import sqlite3
from pathlib import Path
from typing import Any

from werkzeug.utils import secure_filename

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

try:
    import openpyxl
except Exception:
    openpyxl = None

try:
    import xlrd
except Exception:
    xlrd = None

TEXT_ATTACHMENT_EXTS = {
    ".txt",
    ".md",
    ".markdown",
    ".json",
    ".csv",
    ".tsv",
    ".xml",
    ".html",
    ".htm",
    ".yaml",
    ".yml",
    ".ini",
    ".cfg",
    ".conf",
    ".py",
    ".js",
    ".ts",
    ".java",
    ".go",
    ".rb",
    ".php",
    ".sql",
    ".log",
    ".sh",
    ".c",
    ".cpp",
    ".h",
    ".hpp",
}
IMAGE_ATTACHMENT_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}
WORD_ATTACHMENT_EXTS = {".doc", ".docx"}
EXCEL_ATTACHMENT_EXTS = {".xls", ".xlsx"}
DEFAULT_ALLOWED_ATTACHMENT_EXTS = (
    TEXT_ATTACHMENT_EXTS | IMAGE_ATTACHMENT_EXTS | WORD_ATTACHMENT_EXTS | EXCEL_ATTACHMENT_EXTS
)


def file_ext(file_name: str) -> str:
    return Path(file_name).suffix.lower()


def parse_allowed_attachment_exts(raw_exts: str) -> set[str]:
    parsed = set()
    for token in raw_exts.split(","):
        ext = token.strip().lower()
        if not ext:
            continue
        if not ext.startswith("."):
            ext = f".{ext}"
        parsed.add(ext)
    return parsed or set(DEFAULT_ALLOWED_ATTACHMENT_EXTS)


def normalize_uploaded_file_name(raw_name: str, fallback_stem: str) -> str:
    original = (raw_name or "").strip()
    original_ext = file_ext(original)
    secured = secure_filename(original)

    if not secured:
        return f"{fallback_stem}{original_ext}" if original_ext else fallback_stem

    secured_ext = file_ext(secured)
    if secured_ext:
        return secured

    # Some non-ASCII names may collapse to "docx"/"xlsx" without dot.
    if original_ext and secured.lower() == original_ext.lstrip("."):
        return f"{fallback_stem}{original_ext}"

    return f"{secured}{original_ext}" if original_ext else secured


def is_image_mime(mime_type: str) -> bool:
    return mime_type.lower().startswith("image/")


def is_image_attachment(file_name: str, mime_type: str) -> bool:
    return is_image_mime(mime_type) or file_ext(file_name) in IMAGE_ATTACHMENT_EXTS


def is_word_attachment(file_name: str) -> bool:
    return file_ext(file_name) in WORD_ATTACHMENT_EXTS


def is_excel_attachment(file_name: str) -> bool:
    return file_ext(file_name) in EXCEL_ATTACHMENT_EXTS


def is_mime_compatible(ext: str, mime_type: str) -> bool:
    mime = mime_type.strip().lower()
    if not mime or mime == "application/octet-stream":
        return True

    mime_by_ext = {
        ".doc": {"application/msword"},
        ".docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
        ".xls": {"application/vnd.ms-excel"},
        ".xlsx": {"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"},
        ".json": {"application/json"},
        ".csv": {"text/csv", "application/csv", "application/vnd.ms-excel"},
        ".tsv": {"text/tab-separated-values"},
    }
    if ext in IMAGE_ATTACHMENT_EXTS:
        return mime.startswith("image/")
    if ext in TEXT_ATTACHMENT_EXTS:
        expected = mime_by_ext.get(ext)
        return mime.startswith("text/") or (expected is not None and mime in expected)

    expected = mime_by_ext.get(ext)
    if expected is None:
        return True
    return mime in expected


def validate_attachment(file_name: str, mime_type: str, allowed_exts: set[str]) -> tuple[bool, str | None]:
    ext = file_ext(file_name)
    if not ext or ext not in allowed_exts:
        supported = ",".join(sorted(allowed_exts))
        return False, f"不支持的文件类型: {file_name}。仅允许: {supported}"
    if not is_mime_compatible(ext, mime_type):
        return False, f"文件类型与 MIME 不匹配: {file_name} ({mime_type})"
    return True, None


def is_text_attachment(file_name: str, mime_type: str) -> bool:
    mime = mime_type.lower()
    if mime.startswith("text/"):
        return True

    return file_ext(file_name) in TEXT_ATTACHMENT_EXTS


def decode_text_bytes(raw: bytes) -> str:
    for encoding in ("utf-8", "gb18030"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def build_file_text_block(file_name: str, text: str, max_chars: int) -> str:
    cleaned = text.strip()
    truncated = cleaned[:max_chars]
    suffix = "\n[内容已截断]" if len(cleaned) > max_chars else ""
    return f"[文件: {file_name}]\n{truncated}{suffix}\n[文件结束]"


def normalize_extracted_text(text: str) -> str:
    cleaned = text.replace("\r\n", "\n").replace("\r", "\n")
    cleaned = re.sub(r"[^\S\n]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def extract_docx_text(raw: bytes) -> str:
    if DocxDocument is None:
        raise RuntimeError("缺少 python-docx 依赖，无法解析 .docx")
    doc = DocxDocument(io.BytesIO(raw))
    lines: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return normalize_extracted_text("\n".join(lines))


def extract_doc_text(raw: bytes) -> str:
    # Legacy .doc is binary. This is a best-effort fallback extractor.
    utf16_text = raw.decode("utf-16le", errors="ignore")
    utf8_text = raw.decode("utf-8", errors="ignore")
    latin_text = raw.decode("latin1", errors="ignore")
    combined = "\n".join([utf16_text, utf8_text, latin_text])
    candidates = re.findall(r"[A-Za-z0-9\u4e00-\u9fff][A-Za-z0-9\u4e00-\u9fff \t,.;:()_\-/%]{3,}", combined)
    extracted = normalize_extracted_text("\n".join(candidates))
    if extracted:
        return extracted
    return "未能稳定解析 .doc 文档，请优先转换为 .docx 再上传。"


def extract_xlsx_text(raw: bytes) -> str:
    if openpyxl is None:
        raise RuntimeError("缺少 openpyxl 依赖，无法解析 .xlsx")
    wb = openpyxl.load_workbook(io.BytesIO(raw), data_only=True, read_only=True)
    lines: list[str] = []
    for ws in wb.worksheets:
        lines.append(f"[Sheet] {ws.title}")
        for row in ws.iter_rows(values_only=True):
            values = []
            for cell in row:
                if cell is None:
                    values.append("")
                else:
                    values.append(str(cell).strip())
            if any(values):
                lines.append(" | ".join(values))
    return normalize_extracted_text("\n".join(lines))


def extract_xls_text(raw: bytes) -> str:
    if xlrd is None:
        raise RuntimeError("缺少 xlrd 依赖，无法解析 .xls")
    book = xlrd.open_workbook(file_contents=raw)
    lines: list[str] = []
    for sheet in book.sheets():
        lines.append(f"[Sheet] {sheet.name}")
        for row_idx in range(sheet.nrows):
            values = []
            for col_idx in range(sheet.ncols):
                cell_value = sheet.cell_value(row_idx, col_idx)
                values.append(str(cell_value).strip())
            if any(values):
                lines.append(" | ".join(values))
    return normalize_extracted_text("\n".join(lines))


def extract_document_text(file_name: str, raw: bytes) -> str:
    ext = file_ext(file_name)
    if ext == ".docx":
        return extract_docx_text(raw)
    if ext == ".doc":
        return extract_doc_text(raw)
    if ext == ".xlsx":
        return extract_xlsx_text(raw)
    if ext == ".xls":
        return extract_xls_text(raw)
    raise RuntimeError(f"不支持的文档类型: {ext}")


def to_data_url(raw: bytes, mime_type: str) -> str:
    encoded = base64.b64encode(raw).decode("ascii")
    return f"data:{mime_type};base64,{encoded}"


def infer_mime_type(file_name: str, fallback: str = "application/octet-stream") -> str:
    guessed, _ = mimetypes.guess_type(file_name)
    return guessed or fallback


def load_message_attachments(conn: sqlite3.Connection, message_id: int) -> list[sqlite3.Row]:
    return conn.execute(
        """
        SELECT id, file_name, file_path, mime_type, kind, parsed_text, created_at
        FROM message_attachments
        WHERE message_id = ?
        ORDER BY id ASC
        """,
        (message_id,),
    ).fetchall()


def strip_attachment_marker_lines(text: str) -> str:
    lines = text.splitlines()
    filtered = [line for line in lines if not line.strip().startswith("[附件]")]
    return "\n".join(filtered).strip()


def build_message_content_for_model(
    role: str,
    content: str,
    attachments: list[sqlite3.Row],
    max_text_file_chars: int,
) -> Any:
    if role != "user":
        return content

    parts: list[dict[str, Any]] = []
    base_text = strip_attachment_marker_lines(content)
    if base_text:
        parts.append({"type": "text", "text": base_text})

    for attachment in attachments:
        file_name = str(attachment["file_name"])
        mime_type = str(attachment["mime_type"] or "application/octet-stream")
        kind = str(attachment["kind"])
        file_path = Path(str(attachment["file_path"]))

        if kind == "image":
            if file_path.exists():
                raw = file_path.read_bytes()
                parts.append({"type": "image_url", "image_url": {"url": to_data_url(raw, mime_type)}})
            else:
                parts.append({"type": "text", "text": f"[图片附件缺失: {file_name}]"})
            continue

        parsed_text = str(attachment["parsed_text"] or "").strip()
        if parsed_text:
            parts.append({"type": "text", "text": parsed_text})
            continue

        if file_path.exists() and is_text_attachment(file_name, mime_type):
            raw = file_path.read_bytes()
            text = decode_text_bytes(raw)
            parts.append(
                {
                    "type": "text",
                    "text": build_file_text_block(file_name, text, max_text_file_chars),
                }
            )
        else:
            parts.append({"type": "text", "text": f"[二进制文件未解析: {file_name}]"})

    if not parts:
        return ""
    if len(parts) == 1 and parts[0]["type"] == "text":
        return parts[0]["text"]
    return parts


def build_user_display_content(user_text: str, file_names: list[str]) -> str:
    base = user_text.strip()
    if not file_names:
        return base
    attachment_lines = "\n".join([f"[附件] {name}" for name in file_names])
    return f"{base}\n{attachment_lines}".strip()
